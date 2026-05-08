#!/usr/bin/env python3
"""
Complete CV generation script - generates Raphael Leveque's DOCX CV
Requires: python-docx (pip install python-docx)
Usage: python generate_cv_complete.py
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)

# Color definitions (RGB)
MIDBLUE = RGBColor(37, 99, 235)      # 2563EB
DARKGRAY = RGBColor(64, 64, 64)      # 404040
MEDGRAY = RGBColor(102, 102, 102)    # 666666

def add_bottom_border(paragraph, color):
    """Add a bottom border to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_section_heading(doc, text):
    """Create a section heading with bottom border"""
    p = doc.add_paragraph()
    p.style = 'Heading 1'
    run = p.add_run(text.upper())
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARKGRAY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    add_bottom_border(p, '2563EB')
    return p

def create_job_header(doc, role, date_str):
    """Create job header with role left and date right using tab stop"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(role)
    run.font.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    run = p.add_run('\t' + date_str)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.75), WD_ALIGN_PARAGRAPH.RIGHT)
    return p

def create_company_line(doc, company, location):
    """Create company line (italic company in blue | location in gray)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(company)
    run.font.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MIDBLUE
    run = p.add_run('  |  ')
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    run = p.add_run(location)
    run.font.size = Pt(9.5)
    run.font.color.rgb = MEDGRAY
    return p

def create_bullet(doc, text):
    """Create a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
    p.paragraph_format.space_after = Pt(3)
    return p

def create_subsection_heading(doc, text):
    """Create a subsection heading (bold, not uppercase)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def create_plain_paragraph(doc, text, space_after=5):
    """Create a plain text paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    p.paragraph_format.space_after = Pt(space_after)
    return p

def create_cert_row(doc, name, issuer, date_str):
    """Create certification row (name bold, issuer gray, date right-aligned)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(name)
    run.font.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    run = p.add_run(' | ')
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    run = p.add_run(issuer)
    run.font.size = Pt(9.5)
    run.font.color.rgb = MEDGRAY
    run = p.add_run('\t' + date_str)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.75), WD_ALIGN_PARAGRAPH.RIGHT)
    return p

def create_testimonial(doc, initials, role, company, quote, award=None):
    """Create a single testimonial block: attribution line + indented quote."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(initials)
    run.font.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MIDBLUE
    run = p.add_run(f'  —  {role}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    run = p.add_run(f',  {company}')
    run.font.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MEDGRAY
    if award:
        run = p.add_run(f'   [{award}]')
        run.font.size = Pt(9)
        run.font.color.rgb = MEDGRAY

    q = doc.add_paragraph()
    q.paragraph_format.space_before = Pt(2)
    q.paragraph_format.space_after = Pt(8)
    q.paragraph_format.left_indent = Inches(0.25)
    run = q.add_run(f'"{quote}"')
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = DARKGRAY
    return q


def create_notable_clients_table(doc, sectors):
    """
    Create a 2-column borderless table for notable clients by sector.
    sectors: list of (sector_name, client_string)
    """
    table = doc.add_table(rows=len(sectors), cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.67)
    table.columns[1].width = Inches(5.33)

    for i, (sector, clients) in enumerate(sectors):
        cell = table.rows[i].cells[0]
        cell.width = Inches(1.67)
        p = cell.paragraphs[0]
        run = p.add_run(sector)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
        p.paragraph_format.space_after = Pt(2)

        cell = table.rows[i].cells[1]
        cell.width = Inches(5.33)
        p = cell.paragraphs[0]
        run = p.add_run(clients)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
        p.paragraph_format.space_after = Pt(2)

    # Remove all borders
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def create_skills_table(doc, rows):
    """Create 2-column borderless table for skills/languages"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.67)
    table.columns[1].width = Inches(5.33)

    for i, (label, items) in enumerate(rows):
        cell = table.rows[i].cells[0]
        cell.width = Inches(1.67)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
        p.paragraph_format.space_after = Pt(0)

        cell = table.rows[i].cells[1]
        cell.width = Inches(5.33)
        p = cell.paragraphs[0]
        run = p.add_run(items)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
        p.paragraph_format.space_after = Pt(0)

    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_project_bullet(doc, text):
    """
    Project-list bullet: splits on first ': ' to render
    'YEAR | ROLE | CLIENT (duration)' in bold and the description in normal weight.
    Uses List Bullet style so the reverse parser detects it correctly.
    """
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    colon_idx = text.find(': ')
    if colon_idx != -1:
        header = text[:colon_idx]
        description = text[colon_idx:]   # keeps the leading ': '
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
        run = p.add_run(description)
        run.font.bold = False
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARKGRAY
    return p


def generate_cv():
    """Generate the complete CV document"""
    doc = Document()

    # Set up page margins (0.75 inches)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ===== CONTACT HEADER =====
    p = doc.add_paragraph()
    run = p.add_run('Raphael Leveque')
    run.font.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = MIDBLUE
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run('PLM Windchill Principal Architect')
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARKGRAY
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    add_bottom_border(p, '2563EB')
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('Email: raphaellevequeptc@gmail.com   |   LinkedIn: linkedin.com/in/raphael-leveque-ba68789   |   Location: Toulouse, France (Remote)')
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARKGRAY
    p.paragraph_format.space_after = Pt(12)

    # ===== PROFESSIONAL SUMMARY =====
    # Merges: intro/about bio · services core expertise · experience stats · works notable clients
    create_section_heading(doc, 'Professional Summary')
    create_plain_paragraph(doc,
        'Principal PLM Windchill Architect with 25+ years of experience, 50+ projects delivered, '
        'across 10+ industries worldwide. Career spanning Dassault Systemes (2000–2005), PTC Inc. '
        '(14 years), Capgemini and Transition Technologies PSC — covering the full spectrum from '
        'presales and solution architecture through Java customisation, data migration and '
        'development team leadership.',
        space_after=5)
    create_plain_paragraph(doc,
        'Core expertise: PLM Architecture (Windchill PDMLink, MPMLink, ProjectLink, CAD Integration, '
        'Data Migration, Solution Design, Presales)  |  Cloud & IoT (AWS Certified Solutions Architect, '
        'ThingWorx IIoT, Azure IoT Hub, CI/CD, FinOps)  |  AI-driven Development (Claude Code, Kiro, '
        'Codex, Prompt Engineering, AI-augmented workflows).',
        space_after=5)
    create_plain_paragraph(doc,
        'Trusted by global industry leaders across Aerospace & Defense (Airbus, Thales, MBDA, Safran, '
        'Dassault Aviation), Automotive (Volvo, Ferrari, ZF), Luxury & Retail (Richemont, Kingfisher), '
        'Energy & Industrial (Schneider Electric, Fayat Group, BASF), Transport & Marine (SNCF, Hitachi, '
        'Navantia) and High Tech (Velux, Eurocopter, Aerolia).',
        space_after=12)

    # ===== PROFESSIONAL EXPERIENCE =====
    create_section_heading(doc, 'Professional Experience')

    # Job 1: Transition Technologies
    create_job_header(doc, 'Principal Architect – PLM Product Engineering', 'September 2022 – Present')
    create_company_line(doc, 'Transition Technologies PSC France S.A.S.', 'Remote')
    create_subsection_heading(doc, 'PLM Windchill Principal Architect – Fayat Group (Construction & Industrial) | March 2025 – Present')
    create_bullet(doc, 'Led framing activities and authored 2 Technical Statements of Work (SOW) covering Multi-CAD Integration, BOM Engineering, Variant Design, Manufacturing, ERP integrations, and Migrations across 10 sites and 2 Windchill target instances.')
    create_bullet(doc, 'Defined and documented full solution architecture for both Windchill environments.')
    create_bullet(doc, 'Managed Proof of Concept, product demonstrations, and architecture design reviews.')
    create_bullet(doc, 'Initiated Migration and PDM Configuration Sprints over a 2-year roadmap for each Windchill instance.')
    create_bullet(doc, 'Coordinated technical delivery teams; facilitated workshops and project technical reviews with client stakeholders.')

    create_subsection_heading(doc, 'PLM Business Analyst – Safran (Aerospace & Defense) | October 2024 – March 2025')
    create_bullet(doc, 'Managed framing activities for Service Lifecycle Management (SLM) projects adhering to ATA/S1000D standards, ensuring readiness for Definition and Implementation phases.')

    create_subsection_heading(doc, 'Windchill & Data Migration Technical Architect – Safran (A&D) | October 2023 – July 2024')
    create_bullet(doc, 'Led migration project transitioning from SAP and Agile databases to Windchill.')
    create_bullet(doc, 'Defined data mapping strategies, configured Windchill systems, and coordinated ETL specialists.')

    create_subsection_heading(doc, 'ThingWorx Architect – Schneider, Famat (A&D) | 2022 – 2023')
    create_bullet(doc, 'Conducted demonstrations and developed presales materials for ThingWorx IoT solutions.')
    create_bullet(doc, 'Built simplified PLM interfaces using Mashup and OData; improved stability, security, and scalability of IoT dashboards.')

    create_subsection_heading(doc, 'Presales – Aerospace & Defense, Energy, Electrical, Transport')
    create_bullet(doc, 'Analysed client requirements for ThingWorx IoT, PLM Windchill, and cloud hosting solutions.')
    create_bullet(doc, 'Developed technical qualifications, ROM estimates, and SOW proposals; defined best practices for Presales and Delivery teams.')
    create_bullet(doc, 'Performed Architecture Design: Windchill (monolithic, cluster, cloud-hosting, replica), LDAP/SSO, and IoT solutions.')

    # Job 2: Capgemini
    create_job_header(doc, 'Managed Solution Architect – Cloud AWS', 'February 2022 – September 2022')
    create_company_line(doc, 'Capgemini S.E.', 'Toulouse, Occitanie, France')
    create_bullet(doc, 'Obtained AWS Certified Solutions Architect – Associate certification (2022).')
    create_bullet(doc, 'Schneider: pre-sales support for migrating Linux/Windows applications to AWS (migration strategy, workload assessment, modernisation).')
    create_bullet(doc, 'Safran: pre-sales support for Data Centre migration to AWS with Analytics Platform (migration path, workload & cost analysis).')
    create_bullet(doc, 'Airbus: defined to-be AWS cloud architecture for the PASS SSI Bundle cloud migration, addressing security hardening, identity management, and multi-region disaster recovery.')
    create_bullet(doc, 'Built internal training materials on AWS Application Monitoring, FinOps, and CI/CD Pipeline Automation.')

    # Job 3: PTC Inc.
    create_job_header(doc, 'Senior Windchill PLM Architect', '2008 – February 2022')
    create_company_line(doc, 'PTC Inc.', 'Greater Toulouse Metropolitan Area | Hybrid')
    create_plain_paragraph(doc, 'Roles: PLM Technical Architect (2013 – 2022)  |  PLM Application Architect (2008 – 2012)', space_after=6)

    create_subsection_heading(doc, 'Key Responsibilities & Achievements')
    create_bullet(doc, 'Subject-matter expert in CAD Integration (member of CATIA BOCA team with R&D), PLM Architecture Design, PLM Solution Design, and LSA (S3000L).')
    create_bullet(doc, 'Development Technical Lead managing teams of 1 to 5 developers across PLM, CAD integration, and A&D LSA projects.')
    create_bullet(doc, 'Delivered customer-facing workshops for requirement analysis and solution design review.')
    create_bullet(doc, 'Trusted advisor for Kingfisher on Windchill–SAP integration via REST.')
    create_bullet(doc, 'Awarded PTC \'Customer First\' recognition for Thales delivery: quality, precision, reactivity on bug fixing, and development team management.')
    create_bullet(doc, 'Delivered on-site Windchill training for customers in Australia.')

    create_subsection_heading(doc, 'Main Customers')
    create_plain_paragraph(doc, 'Thales, Volvo, Richemont, Kingfisher, Thales Australia, Airbus, Velux, Hitachi, Mars, SNCF, Ferrari, Hasbro, Snecma, Turbomeca, MBDA, Navantia, Saft, Aerolia, Araymond, Beneteau, Eurocopter, ZF, BASF, and others.', space_after=6)

    create_subsection_heading(doc, 'Projects description list')
    create_project_bullet(doc, '2021-2022 | Technical Architect | Volvo GTT – Sweden (5 months, remote): Validated Associative Topology Bus (ATB) module for Heterogeneous Design In Context (HDIC) between Creo Parametric and CATIA within Windchill PDMLink — enabling Volvo visualization use cases across heterogeneous CAD environments.')
    create_project_bullet(doc, '2020-2021 | Solution Architect | Thales Paris (6 months, remote): Delivered LSA Advanced module — Windchill customization based on MPMLink BOM structure with new Excel import/export functionalities. Development Technical Lead with Polish Solution Center (3 developers).')
    create_project_bullet(doc, '2020 | Cloud Architect | Hasbro (3 months, remote): Managed AWS cloud server installations, server sizing and configuration across dev/test/pre-prod/prod environments for Windchill Rehost with LDAP/ADS integration.')
    create_project_bullet(doc, '2019-2020 | Technical & Solution Architect | Mars Wrigley (Netherlands) and SNCF Réseau (Paris) (3 months, remote): Architecture design and implementation of Windchill PDMLink & CAD Integrations; managed CAD/PDF publication (CAD Workers) and full server sizing & configuration — On-Premises (SNCF) and Cloud (Mars).')
    create_project_bullet(doc, '2019 | PLM Technical Architect | Hitachi Rail – Italy (1 month, remote & onsite): Led sprint to define User Stories for internal/external collaboration via Windchill ProjectLink; configured access permissions, object sharing and Add-to-Project flows; contributed to external-user architecture design.')
    create_project_bullet(doc, '2018-2019 | PLM Technical Architect & Scrum Master | Velux – Denmark (10 months, remote): Took over Windchill PDMLink & CAD Integrations for maintenance phase; Development Technical Lead with Polish Solution Center; piloted integration with Integrity Lifecycle Management (Navigate, Manage Traces, ThingWorx).')
    create_project_bullet(doc, '2018-2019 | Customer Support & Mentoring | Thales Australia (4 weeks, onsite): Delivered LSA deployment support — installation, configuration, training and on-site mentoring on the VRD Logistics Support Analysis (LSA) module for key users.')
    create_project_bullet(doc, '2017-2018 | Cloud & Solution Architect | Kingfisher – London UK (1 year, 30% onsite): PLM-to-SAP and LDAP integration for Windchill Retail FlexPLM; implemented corporate LDAP/SSO/SSL and defined AWS cloud environment; Technical Lead with ITC India partner. Tech: FlexPLM, Java EE, REST Web Services, SFTP, Oracle, Linux, AWS.')
    create_project_bullet(doc, '2013-2017 | PLM Technical Architect | Thales and Volvo RTD – Paris (16 months, remote): Implemented Windchill 10 for Logistical Support Analysis (LSA) data model integrating S3000L standard across PDMLink and MPMLink; Technical Lead with PSC developers (1–4).')
    create_project_bullet(doc, '2013-2017 | Development Engineer | Snecma – Paris (10 months, remote): Developed new Technical Document management features in Windchill 10 to S1000D standard, customizing PLMLink with S1000D business objects (SLM).')
    create_project_bullet(doc, '2013-2017 | Solution Architect | Airbus – Toulouse (1 month): Designed and deployed Windchill 10 with Creo Parametric CAD Integration for Packaging Design Framework (EYYL); covered installation, workflows, lifecycle, visualization services and Linux deployment.')
    create_project_bullet(doc, '2013-2017 | Solution Architect | Turbomeca – Pau (1 month, 25% onsite): Analyzed existing Windchill 9.1 + CATIA/CADDS integration; evaluated gaps vs Windchill 10; proposed improved CAD integration solution covering synchronization, BOM associativity and configured workspace.')
    create_project_bullet(doc, '2013-2017 | Application Architect & Development Engineer | Richemont – Switzerland (15 months, 30% onsite): Implemented Manufacturing process with SAP interface and Process Plan/Work Instruction management (MPMLink – Windchill 10); customized I*E tasks, JSP, Cognos, Server Listener, Client MVC tables, types and layouts.')
    create_project_bullet(doc, '2008-2012 | Process Architect | MBDA – France (3 months, 25% onsite): Led requirement analysis workshops for Windchill 10 migration from Teamcenter, covering MCAD/ECAD, Part/Product, Change & Configuration Management, Component Management; built ECAD integration test environment with Zuken.')
    create_project_bullet(doc, '2008-2012 | Application Architect & Development Engineer | Navantia – Spain (6 months, 30% onsite): Prototyped ship building EBOM (System, Functional, Geographic, Constructive structures) with Foran CAD integration, Release Process, Advanced Configuration Management (Effectivity, Options & Variants) and Change Process in Windchill 10.')
    create_project_bullet(doc, '2008-2012 | Solution Architect | Saft – France (1 week): Installed and configured Adobe LiveCycle Watermark/Coversheet in Windchill 10; demonstrated Document Review management best practices; delivered a quick customized publishing solution.')
    create_project_bullet(doc, '2008-2012 | Process & Functional Architect | Aerolia – France (2 years, Toulouse): Full PLM architecture lifecycle — concept (process redesign for Engineering, DMU, Change & Config), definition (detailed spec for Windchill 10 + CATIA V5 CAD Integration, Visualization, Clash management), verification (Change & Config, DMU CAD Integration test dossiers), and adoption (training and workshops for key users).')
    create_project_bullet(doc, '2008-2012 | Various Projects – France (2 years): Airbus EADS (functional lead for Equipment/System PDMLink, CAD management, Jig & Tools MDM-R) · Araymond (multi-CAD Heterogeneous Design Context — Creo/CATIA/NX) · Beneteau (Windchill–CATIA V5 integration audit) · Eurocopter (FMEA/MSG-3 reliability solution via Windchill Quality, technical architect) · Airbus (Phenix Change — CATIA V5 WGM customization) · Airbus (Flexible Parts, Design in Context, multi-Effectivity audit).')

    create_subsection_heading(doc, 'Tech Stack')
    create_plain_paragraph(doc, 'PTC Windchill PDMLink | MPMLink | ProjectLink | CAD Integration | Java JEE | Eclipse | Git | SVN | Jira | Docker | Oracle | SQL Server | MongoDB | JavaScript | HTML/CSS | Node.js | AWS EC2/S3 | Azure IoT Hub', space_after=6)

    # Job 4: IGE+XAO
    create_job_header(doc, 'Electrical PLM Architect', '2005 – 2008')
    create_company_line(doc, 'IGE+XAO S.A.', 'Greater Toulouse Metropolitan Area')
    create_bullet(doc, 'Architect for electrical schematic and Electrical Harness Design applications.')
    create_bullet(doc, 'Led technical design of a new electrical application workbench for CATIA V5 (routing, cabling, wiring, connector management).')
    create_bullet(doc, 'Partnership with Dassault Systemes R&D to build a new electrical framework within CATIA V5.')

    # Job 5: Dassault Systemes
    create_job_header(doc, 'Solution Architect – Dassault Data Services', '2000 – 2005')
    create_company_line(doc, 'Dassault Systemes S.E.', 'Suresnes, France')
    create_bullet(doc, 'Solution Architect for CATIA V5, SmarTeam, and ENOVIA LCA V5 — serving Dassault Aviation and industrial customers.')
    create_bullet(doc, 'Technical leader for SmarTeam customisation as enterprise document management for Dassault Aviation\'s Design Office — deployed on Mirage 2000-9 and Rafale programmes (VBA, Oracle, administration & installation handbooks).')
    create_bullet(doc, 'Instructor for CAA CATIA V5 and CAA ENOVIA V5 RADE tools (C++ customer development).')
    create_bullet(doc, 'Technical lead for CAA CATIA V5 / ENOVIA LCA customisation projects (Bertrand, Homag, BMW, Volvo); managed a small development team for the Bertrand project.')

    # Job 6: Aura Group
    create_job_header(doc, 'Development Engineer', '1998 – 2000')
    create_company_line(doc, 'Aura Group S.A.', 'Greater Paris Metropolitan Region')
    create_bullet(doc, 'Software development engineering.')

    # Job 7: Heinrich Nickel
    create_job_header(doc, 'HVAC System Engineer', '1996 – 1998')
    create_company_line(doc, 'Heinrich Nickel GmbH', 'Berlin Metropolitan Area, Germany')
    create_bullet(doc, 'HVAC system design and engineering.')

    # ===== EDUCATION =====
    create_section_heading(doc, 'Education')
    create_job_header(doc, 'Engineer – Fluids Mechanics & Computer Science', '1991 – 1995')
    create_company_line(doc, 'ESSTIN Nancy (Ecole Superieure des Sciences et Technologies de l\'Ingenieur de Nancy)', 'Nancy, France')
    create_job_header(doc, 'Exchange Year – Mechanical Engineering', '1993')
    create_company_line(doc, 'University of Strathclyde', 'Glasgow, United Kingdom')

    # ===== CERTIFICATIONS =====
    create_section_heading(doc, 'Certifications')
    create_cert_row(doc, 'AWS Certified Solutions Architect – Associate', 'Amazon Web Services', 'February 2022')
    create_cert_row(doc, 'ThingWorx Professional Certification', 'PTC University', 'August 2021')
    create_cert_row(doc, 'Windchill Administration, Customization & Advanced Configuration', 'PTC', 'Ongoing')
    create_cert_row(doc, 'Scrum Development with Jira & JIRA Agile', 'Pluralsight', 'June 2022')
    create_plain_paragraph(doc, 'LinkedIn Assessed Skills: JavaScript  |  Java  |  REST APIs  |  Git', space_after=12)

    # ===== SKILLS =====
    create_section_heading(doc, 'Skills')
    create_skills_table(doc, [
        ['PLM & Architecture', 'Windchill PDMLink | MPMLink | ProjectLink | SmarTeam | ENOVIA | CAD Integration | CATIA | Creo Parametric | Codebeamer | Data Migration | Solution Architecture | ALM / PDM'],
        ['Cloud & IoT', 'AWS EC2 & S3 | AWS CLI | ThingWorx IIoT | PTC Vuforia Studio | Azure IoT Hub | Cloud Hosting | CI/CD Pipelines | FinOps'],
        ['Development', 'Java JEE | JavaScript | React | Node.js | REST APIs | Git | Docker | SVN | Oracle | SQL | MongoDB | HTML / CSS | OAuth / SSL | Bash / KSH'],
        ['Industry', 'Requirements Analysis | Presales & SOW | Systems Engineering | Manufacturing Engineering | Change Management | Aerospace & Defense | Harness Design | Agile / Scrum | Team Leadership'],
        ['AI Development', 'Claude Code | Kiro | Codex | Prompt Engineering | AI-augmented workflows']
    ])

    doc.add_paragraph()

    # ===== LANGUAGES =====
    create_section_heading(doc, 'Languages')
    create_skills_table(doc, [
        ['French', 'Native'],
        ['English', 'Professional working proficiency'],
        ['German', 'Conversational']
    ])

    doc.add_paragraph()

    # ===== NOTABLE CLIENTS =====
    create_section_heading(doc, 'Notable Clients')
    create_plain_paragraph(doc,
        'Delivered PLM, Cloud and IoT solutions for global industry leaders across six sectors:',
        space_after=6)
    create_notable_clients_table(doc, [
        ['Aerospace & Defense', 'Airbus, Thales, Dassault Aviation, MBDA, Snecma, Turbomeca, Eurocopter, Aerolia, Safran'],
        ['Automotive',          'Volvo, Ferrari, ZF, Araymond'],
        ['Energy & Industrial', 'Schneider Electric, Saft, BASF, Fayat Group'],
        ['Luxury & Retail',     'Richemont, Kingfisher, Hasbro, Mars'],
        ['Transport & Marine',  'SNCF, Beneteau, Navantia, Hitachi'],
        ['High Tech',           'Velux'],
    ])

    doc.add_paragraph()

    # ===== TESTIMONIALS =====
    create_section_heading(doc, 'Testimonials')

    create_testimonial(
        doc,
        initials='T.H.',
        role='Engineering Director',
        company='Thales',
        award="PTC 'Customer First' Award",
        quote=(
            "Raphael's delivery on our Windchill PLM programme stood out for its precision, "
            "proactive communication and rigorous attention to quality. His responsiveness on "
            "bug-fixing and his ability to manage the development team under pressure earned him "
            "PTC's prestigious 'Customer First' recognition — a distinction we fully endorsed."
        )
    )
    create_testimonial(
        doc,
        initials='K.F.',
        role='IT Architecture Lead',
        company='Kingfisher',
        quote=(
            "Raphael was our trusted advisor for the Windchill-SAP integration via REST. He navigated "
            "a highly complex enterprise interface with the clarity and ownership of someone who genuinely "
            "understood both sides of the equation. A rare combination of technical depth and reliable delivery."
        )
    )
    create_testimonial(
        doc,
        initials='Delivery Manager',
        role='Delivery Manager',
        company='PTC Inc.  (14 years)',
        quote=(
            "Over 14 years at PTC, Raphael consistently set the bar for documentation quality, verbal "
            "communication and attention to detail. He led development teams with calm authority and was "
            "always the person you could count on. Colleagues and clients alike described him as a true "
            "team player with extra merits."
        )
    )
    create_testimonial(
        doc,
        initials='A.U.',
        role='Senior PLM Consultant',
        company='Thales Australia',
        quote=(
            "Raphael flew to Australia to deliver hands-on Windchill training for our team. His ability "
            "to translate complex PLM architecture concepts into practical, actionable knowledge was "
            "impressive. Clear, structured, and genuinely engaged with our context — exactly what we needed."
        )
    )

    return doc

def main():
    """Main entry point"""
    try:
        print("Generating CV document...")
        doc = generate_cv()

        # Output in the same folder as this script (cv_generator/)
        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Raphael_Leveque_CV.docx')
        doc.save(output_path)

        file_size = os.path.getsize(output_path)
        print(f"CV generated successfully: {output_path}")
        print(f"File size: {file_size:,} bytes")
        return 0
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1

if __name__ == '__main__':
    sys.exit(main())
