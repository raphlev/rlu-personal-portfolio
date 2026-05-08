#!/usr/bin/env python3
"""
cv_to_portfolio.py
------------------
Reads Raphael_Leveque_CV.docx and merges any changes back into
public/portfolio.md.

Sections synced from CV → portfolio.md:
  timeline        ← Professional Experience jobs
  skills          ← Skills table + Certifications
  about           ← Education block only (bio left untouched)
  works           ← Notable Clients table
  testimonials    ← Testimonials section

Sections intentionally NOT touched (no equivalent in the CV):
  intro, navbar, services, experience stats, portfolio cards, footer

Usage:
    python cv_to_portfolio.py
    python cv_to_portfolio.py --cv path/to/cv.docx --portfolio path/to/portfolio.md

Requires:  pip install python-docx
"""

import sys
import os
import re
import json
import argparse
from copy import deepcopy

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Run:  pip install python-docx")
    sys.exit(1)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def para_text(p):
    """Return full text of a paragraph, stripping tab-separated right-aligned date."""
    return p.text.strip()

def is_section_heading(p):
    """Section headings are uppercase, bold, short (< 60 chars), and have a bottom border."""
    txt = para_text(p)
    if not txt or len(txt) > 60:
        return False
    # All-caps check
    if txt != txt.upper():
        return False
    # Check for bottom border in paragraph XML
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            bottom = pBdr.find(qn('w:bottom'))
            if bottom is not None:
                return True
    return False

def is_job_header(p):
    """Job headers have a right-aligned tab stop and bold first run."""
    if not p.runs:
        return False
    first_bold = p.runs[0].bold
    has_tab = '\t' in p.text
    return bool(first_bold and has_tab)

def is_company_line(p):
    """Company lines have an italic first run."""
    if not p.runs:
        return False
    return bool(p.runs[0].italic)

def is_bullet(p):
    """Detect List Bullet style."""
    return p.style.name.startswith('List Bullet') if p.style else False

def is_subsection_heading(p):
    """Bold non-uppercase short paragraph = subsection heading."""
    txt = para_text(p)
    if not txt or len(txt) > 120:
        return False
    if p.runs and p.runs[0].bold and txt != txt.upper():
        return True
    return False

def split_tab_line(text):
    """Split 'Left content\\tRight content' into (left, right)."""
    parts = text.split('\t', 1)
    return (parts[0].strip(), parts[1].strip()) if len(parts) == 2 else (parts[0].strip(), '')

def month_abbrev_to_full(date_str):
    """'Sep 2022' → 'September 2022' etc. No-op if already full."""
    abbrevs = {
        'Jan': 'January', 'Feb': 'February', 'Mar': 'March', 'Apr': 'April',
        'May': 'May', 'Jun': 'June', 'Jul': 'July', 'Aug': 'August',
        'Sep': 'September', 'Oct': 'October', 'Nov': 'November', 'Dec': 'December'
    }
    for short, full in abbrevs.items():
        date_str = re.sub(r'\b' + short + r'\b', full, date_str)
    return date_str

def strip_company_suffix(name):
    """Remove trailing legal suffix for matching: 'PTC Inc.' → 'PTC'."""
    return re.sub(
        r'\s+(S\.A\.S\.|S\.E\.|Inc\.|GmbH|S\.A\.|Ltd\.|LLC|PSC|plc|N\.V\.|AG|S\.p\.A\.)\.?$',
        '', name, flags=re.IGNORECASE
    ).strip()


# ─── Parse DOCX ───────────────────────────────────────────────────────────────

def parse_cv(docx_path):
    """
    Parse the CV document and return a structured dict with the data
    that maps back to portfolio.md sections.
    """
    doc = Document(docx_path)
    paras = [p for p in doc.paragraphs]

    result = {
        'name': '',
        'title': '',
        'email': '',
        'linkedin': '',
        'location': '',
        'summary_paras': [],  # list of paragraphs (summary is now multi-line)
        'jobs': [],           # list of job dicts
        'education': [],      # list of education dicts
        'certifications': [], # list of cert dicts
        'skills': {},         # {label: [items]}
        'languages': {},      # {label: level}
        'notable_clients': [], # list of (sector, clients_string)
        'testimonials': [],   # list of testimonial dicts
    }

    current_section = None
    current_job = None
    current_sub = None
    pending_testimonial = None   # attribution line waiting for its quote
    i = 0

    while i < len(paras):
        p = paras[i]
        txt = para_text(p)

        # ── Header block (first ~5 paragraphs) ────────────────────────────────
        if not result['name'] and txt and not is_section_heading(p):
            # Name is first non-empty paragraph
            result['name'] = txt
            i += 1
            continue

        if not result['title'] and txt and not result['summary'] and not is_section_heading(p) and not is_company_line(p):
            if result['name'] and not result['title']:
                result['title'] = txt
                i += 1
                continue

        # Contact line: "Email: x   |   LinkedIn: y   |   Location: z"
        if 'Email:' in txt and 'LinkedIn:' in txt:
            email_m = re.search(r'Email:\s*(\S+)', txt)
            linkedin_m = re.search(r'LinkedIn:\s*(\S+)', txt)
            location_m = re.search(r'Location:\s*(.+?)(?:\s*\||\s*$)', txt)
            if email_m:
                result['email'] = email_m.group(1)
            if linkedin_m:
                result['linkedin'] = linkedin_m.group(1)
            if location_m:
                result['location'] = location_m.group(1).strip()
            i += 1
            continue

        # ── Section headings ──────────────────────────────────────────────────
        if is_section_heading(p):
            current_section = txt.strip().upper()
            current_job = None
            current_sub = None
            i += 1
            continue

        # ── PROFESSIONAL SUMMARY (multi-paragraph) ───────────────────────────
        if current_section == 'PROFESSIONAL SUMMARY':
            if txt:
                result['summary_paras'].append(txt)
            i += 1
            continue

        # ── PROFESSIONAL EXPERIENCE ──────────────────────────────────────────
        if current_section == 'PROFESSIONAL EXPERIENCE':
            if is_job_header(p):
                role, date_range = split_tab_line(txt)
                date_range = month_abbrev_to_full(date_range)
                current_job = {
                    'role': role,
                    'date_range': date_range,
                    'company': '',
                    'location': '',
                    'subsections': [],
                    'bullets': [],
                    'plain': [],
                }
                current_sub = None
                result['jobs'].append(current_job)
            elif current_job and is_company_line(p):
                # "Company S.A.S.  |  Location"
                parts = txt.split('|', 1)
                current_job['company'] = parts[0].strip()
                if len(parts) > 1:
                    current_job['location'] = parts[1].strip()
            elif current_job and is_subsection_heading(p):
                current_sub = {'title': txt, 'items': []}
                current_job['subsections'].append(current_sub)
            elif current_job and is_bullet(p):
                if current_sub:
                    current_sub['items'].append(txt)
                else:
                    current_job['bullets'].append(txt)
            elif current_job and txt and not is_section_heading(p):
                if current_sub:
                    current_sub['items'].append(txt)
                else:
                    current_job['plain'].append(txt)
            i += 1
            continue

        # ── EDUCATION ────────────────────────────────────────────────────────
        if current_section == 'EDUCATION':
            if is_job_header(p):
                degree, years = split_tab_line(txt)
                edu = {'degree': degree, 'years': years, 'school': '', 'location': ''}
                result['education'].append(edu)
            elif result['education'] and is_company_line(p):
                parts = txt.split('|', 1)
                result['education'][-1]['school'] = parts[0].strip()
                if len(parts) > 1:
                    result['education'][-1]['location'] = parts[1].strip()
            i += 1
            continue

        # ── CERTIFICATIONS ───────────────────────────────────────────────────
        if current_section == 'CERTIFICATIONS':
            if txt and '\t' in txt:
                # "Name | Issuer \t Date"
                left, date = split_tab_line(txt)
                name_parts = left.split('|', 1)
                cert_name = name_parts[0].strip()
                issuer = name_parts[1].strip() if len(name_parts) > 1 else ''
                result['certifications'].append({
                    'name': cert_name,
                    'issuer': issuer,
                    'date': date,
                })
            i += 1
            continue

        # ── SKILLS (table) ───────────────────────────────────────────────────
        if current_section == 'SKILLS':
            # Tables are processed separately below
            i += 1
            continue

        # ── LANGUAGES (table) ────────────────────────────────────────────────
        if current_section == 'LANGUAGES':
            i += 1
            continue

        # ── NOTABLE CLIENTS (table) ──────────────────────────────────────────
        if current_section == 'NOTABLE CLIENTS':
            # Table is processed separately below
            i += 1
            continue

        # ── TESTIMONIALS ─────────────────────────────────────────────────────
        if current_section == 'TESTIMONIALS':
            if not txt:
                i += 1
                continue

            # Attribution line: "T.H.  —  Role,  Company   [Award]"
            # Detected by: bold first run and contains em-dash
            if p.runs and p.runs[0].bold and ('—' in txt or '--' in txt):
                # Parse: initials  —  role,  company  [award]
                dash_sep = '—' if '—' in txt else '--'
                parts = txt.split(dash_sep, 1)
                initials = parts[0].strip().rstrip('.')
                # Remove dots to get clean initials (T.H. → TH)
                initials_clean = initials.replace('.', '')
                rest = parts[1].strip() if len(parts) > 1 else ''
                award = ''
                award_m = re.search(r'\[([^\]]+)\]', rest)
                if award_m:
                    award = award_m.group(1).strip()
                    rest = rest[:award_m.start()].strip()
                role_company = [x.strip() for x in rest.split(',', 1)]
                role = role_company[0]
                company = role_company[1] if len(role_company) > 1 else ''
                # Strip legal suffixes and duration info from company (e.g. "PTC Inc.  (14 years)" → "PTC")
                company = re.sub(r'\s*\(.*?\)', '', company).strip()
                company = re.sub(
                    r'\s+(S\.A\.S\.|S\.E\.|Inc\.|GmbH|S\.A\.|Ltd\.|LLC)\.?$',
                    '', company, flags=re.IGNORECASE
                ).strip()
                pending_testimonial = {
                    'initials': initials_clean,
                    'role': role,
                    'company': company,
                    'award': award,
                    'quote': '',
                }
            # Quote line: italic paragraph starting with "  (set by indented italic style)
            elif pending_testimonial is not None:
                # Strip surrounding quotes added by the generator
                quote = txt.strip('"').strip()
                pending_testimonial['quote'] = quote
                result['testimonials'].append(pending_testimonial)
                pending_testimonial = None

            i += 1
            continue

        i += 1

    # ── Parse tables (Skills, Languages, Notable Clients) ────────────────────
    lang_labels      = {'French', 'English', 'German', 'Spanish', 'Italian', 'Portuguese'}
    sector_labels    = {'Aerospace & Defense', 'Automotive', 'Energy & Industrial',
                        'Luxury & Retail', 'Transport & Marine', 'High Tech',
                        'Energy', 'Industrial', 'Marine', 'Transport'}

    for table in doc.tables:
        if len(table.columns) != 2:
            continue
        rows_data = []
        for row in table.rows:
            cells = row.cells
            label = cells[0].text.strip()
            items_text = cells[1].text.strip()
            if label:
                rows_data.append((label, items_text))

        if not rows_data:
            continue

        first_label = rows_data[0][0]

        if first_label in lang_labels:
            # Languages table
            for label, value in rows_data:
                result['languages'][label] = value
        elif first_label in sector_labels or 'Aerospace' in first_label:
            # Notable Clients table
            result['notable_clients'] = [(label, items) for label, items in rows_data]
        else:
            # Skills table
            for label, items_text in rows_data:
                items = [x.strip() for x in items_text.split('|') if x.strip()]
                result['skills'][label] = items

    return result


# ─── Map parsed CV → portfolio.md sections ───────────────────────────────────

def build_timeline_json(jobs):
    """Convert parsed jobs list to the timeline JSON block used in portfolio.md."""
    entries = []
    for job in jobs:
        entry = {
            "period": job['date_range'],
            "company": job['company'],
            "role": job['role'],
            "location": job['location'],
            "tags": [],
        }
        if job['subsections']:
            entry['subsections'] = []
            for sub in job['subsections']:
                entry['subsections'].append({
                    "title": sub['title'],
                    "items": sub['items'],
                })
        elif job['bullets']:
            entry['highlights'] = job['bullets']

        entries.append(entry)
    return entries


def build_skills_markdown(skills_dict):
    """Rebuild the ## skills section skill groups from parsed CV."""
    lines = []
    accent_map = {
        'PLM & Architecture': '#DDF8FE',
        'Cloud & IoT': '#fff3d4',
        'Development': '#f0e6ff',
        'Industry Knowledge': '#e6f4ea',
        'Industry': '#e6f4ea',
        'AI Development': '#fde8d8',
    }
    for label, items in skills_dict.items():
        lines.append(f'\n### {label}')
        accent = accent_map.get(label, '#f0f0f0')
        lines.append(f'**accent:** {accent}')
        for item in items:
            # Strip unicode check-marks that may have been added
            clean = re.sub(r'\s*[✓✔]\s*', '', item).strip()
            lines.append(f'- {clean}')
    return '\n'.join(lines)


def build_certs_json(certs):
    """Rebuild certifications JSON block."""
    cert_meta = {
        'AWS Certified Solutions Architect': {
            'accent': '#FF9900', 'icon': '☁️'
        },
        'ThingWorx': {
            'accent': '#0066CC', 'icon': '🔗'
        },
        'Windchill': {
            'accent': '#00A9E0', 'icon': '⚙️'
        },
        'Scrum': {
            'accent': '#F15B2A', 'icon': '🔄'
        },
    }
    result = []
    for c in certs:
        meta = {'accent': '#888888', 'icon': '📜'}
        for key, val in cert_meta.items():
            if key.lower() in c['name'].lower():
                meta = val
                break
        result.append({
            'name': c['name'],
            'issuer': c['issuer'],
            'date': c['date'],
            'accent': meta['accent'],
            'icon': meta['icon'],
        })
    return result


def build_works_markdown(notable_clients):
    """Rebuild ## works section from Notable Clients table rows."""
    lines = [
        '\n**heading1:** Trusted by Global',
        '**heading2:** Industry Leaders',
        '**description:** Delivering PLM excellence for Airbus, Thales, Safran, Volvo, Richemont, '
        'Ferrari, SNCF, Schneider, MBDA and many more — across Aerospace & Defense, Automotive, '
        'Retail, Energy, Shipbuilding and beyond.',
        '**ctaLabel:** Work with Me',
    ]
    for sector, clients in notable_clients:
        lines.append(f'\n### {sector}')
        # Convert comma-separated clients to middle-dot separated (portfolio.md format)
        client_list = [c.strip() for c in clients.split(',') if c.strip()]
        lines.append(' · '.join(client_list))
    return '\n'.join(lines) + '\n'


def build_testimonials_markdown(testimonials):
    """Rebuild ## testimonials section from parsed testimonial dicts."""
    lines = [
        '\n**heading1:** Colleagues & clients',
        '**heading2:** value precision,',
        '**heading3:** reliability & partnership',
        '\nFormat for subsection names: `INITIALS | Role | Company`',
    ]
    for t in testimonials:
        initials = t['initials']
        role     = t['role']
        company  = t['company']
        award    = t.get('award', '')
        quote    = t['quote']
        lines.append(f'\n### {initials} | {role} | {company}')
        if award:
            lines.append(f'**award:** {award}')
        lines.append('')
        lines.append(quote)
    return '\n'.join(lines) + '\n'


# ─── Read / patch / write portfolio.md ───────────────────────────────────────

def read_portfolio(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def replace_section_block(content, section_name, new_block):
    """
    Replace everything between '## <section_name>' and the next '---' separator
    (or end of file) with new_block.
    """
    pattern = re.compile(
        r'(## ' + re.escape(section_name) + r'\n)(.*?)(\n---|\Z)',
        re.DOTALL
    )
    replacement = r'\g<1>' + new_block + r'\g<3>'
    new_content, n = pattern.subn(replacement, content)
    if n == 0:
        print(f"  WARNING: section '## {section_name}' not found in portfolio.md")
    return new_content


def patch_portfolio(content, cv):
    """Apply CV changes to portfolio.md content string. Returns updated content."""

    # ── timeline ──────────────────────────────────────────────────────────────
    timeline_entries = build_timeline_json(cv['jobs'])
    timeline_json = json.dumps(timeline_entries, indent=2, ensure_ascii=False)
    timeline_block = (
        '\n**heading1:** Professional\n'
        '**heading2:** Career History\n\n'
        '```json\n' + timeline_json + '\n```\n'
    )
    content = replace_section_block(content, 'timeline', timeline_block)

    # ── skills ────────────────────────────────────────────────────────────────
    certs = build_certs_json(cv['certifications'])
    certs_json = json.dumps(certs, indent=2, ensure_ascii=False)
    skills_groups = build_skills_markdown(cv['skills'])
    # Preserve existing heading/assessments lines
    skills_block = (
        '\n**heading1:** Skills &\n'
        '**heading2:** Certifications\n'
        '**assessments:** JavaScript, Java, REST APIs, Git\n'
        + skills_groups + '\n\n'
        '```json\n' + certs_json + '\n```\n'
    )
    content = replace_section_block(content, 'skills', skills_block)

    # ── about — education only ────────────────────────────────────────────────
    edu_lines = []
    for edu in cv['education']:
        flag = '🇫🇷' if 'France' in edu.get('location', '') or 'Nancy' in edu.get('school', '') else \
               '🇬🇧' if 'UK' in edu.get('location', '') or 'United Kingdom' in edu.get('location', '') else '🎓'
        edu_lines.append(f'- {flag} | {edu["years"]} | {edu["school"]} | {edu["degree"]}')

    edu_pattern = re.compile(
        r'(### education\n\nFormat:.*?\n\n)((?:- .+\n?)+)',
        re.DOTALL
    )
    new_edu_list = '\n'.join(edu_lines) + '\n'
    content = edu_pattern.sub(r'\g<1>' + new_edu_list, content)

    # ── works (notable clients) ───────────────────────────────────────────────
    if cv['notable_clients']:
        works_block = build_works_markdown(cv['notable_clients'])
        content = replace_section_block(content, 'works', works_block)
    else:
        print("  INFO: No Notable Clients table found in CV — ## works left unchanged.")

    # ── testimonials ──────────────────────────────────────────────────────────
    if cv['testimonials']:
        testimonials_block = build_testimonials_markdown(cv['testimonials'])
        content = replace_section_block(content, 'testimonials', testimonials_block)
    else:
        print("  INFO: No testimonials found in CV — ## testimonials left unchanged.")

    return content


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Sync CV changes back to portfolio.md')
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # CV lives alongside this script in cv_generator/
    # portfolio.md is one level up in public/
    repo_root = os.path.normpath(os.path.join(script_dir, '..'))
    parser.add_argument(
        '--cv',
        default=os.path.join(script_dir, 'Raphael_Leveque_CV.docx'),
        help='Path to the Word CV (default: cv_generator/Raphael_Leveque_CV.docx)'
    )
    parser.add_argument(
        '--portfolio',
        default=os.path.join(repo_root, 'public', 'portfolio.md'),
        help='Path to portfolio.md (default: <repo_root>/public/portfolio.md)'
    )
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Print diff summary without writing changes'
    )
    args = parser.parse_args()

    if not os.path.exists(args.cv):
        print(f"ERROR: CV file not found: {args.cv}")
        sys.exit(1)
    if not os.path.exists(args.portfolio):
        print(f"ERROR: portfolio.md not found: {args.portfolio}")
        sys.exit(1)

    print(f"Reading CV: {args.cv}")
    cv = parse_cv(args.cv)

    print(f"\nParsed:")
    print(f"  Name           : {cv['name']}")
    print(f"  Summary paras  : {len(cv['summary_paras'])}")
    print(f"  Jobs           : {len(cv['jobs'])}")
    print(f"  Education      : {len(cv['education'])}")
    print(f"  Certifications : {len(cv['certifications'])}")
    print(f"  Skill groups   : {list(cv['skills'].keys())}")
    print(f"  Languages      : {list(cv['languages'].keys())}")
    print(f"  Notable clients: {len(cv['notable_clients'])} sectors")
    print(f"  Testimonials   : {len(cv['testimonials'])}")

    print(f"\nReading portfolio: {args.portfolio}")
    original = read_portfolio(args.portfolio)
    updated = patch_portfolio(original, cv)

    if updated == original:
        print("\nNo changes detected — portfolio.md is already in sync.")
        return

    if args.dry_run:
        print("\n[DRY RUN] Changes would be written. Use without --dry-run to apply.")
        # Show rough diff summary
        orig_lines = original.splitlines()
        upd_lines = updated.splitlines()
        added = sum(1 for l in upd_lines if l not in orig_lines)
        removed = sum(1 for l in orig_lines if l not in upd_lines)
        print(f"  ~{added} lines added, ~{removed} lines removed")
        return

    # Write updated portfolio.md
    with open(args.portfolio, 'w', encoding='utf-8') as f:
        f.write(updated)

    print(f"\nportfolio.md updated successfully.")
    print(f"Sections patched: timeline, skills, about (education)")
    print(f"\nSections NOT modified (manual-only):")
    print(f"  intro, navbar, services, experience, works, portfolio cards,")
    print(f"  testimonials, footer")


if __name__ == '__main__':
    main()
